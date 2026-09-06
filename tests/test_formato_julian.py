"""Tests del formato de Julián: plantilla, perfil y exportación a Word.

Las medidas que se comprueban aquí salen de su documento de estilo:

    Diálogos:  izquierda 0,63 cm con sangría francesa 0,63,
               posterior 18 pt, interlineado mínimo.
    Narrador:  primera línea 1,25 cm, posterior 18 pt, justificado.
"""

from __future__ import annotations

import pytest

from vozclip import exportar_word, perfiles, plantillas
from vozclip.fuentes import ErrorFuente

TEXTO = """     Aquella noche no dormí. El taller olía a papel viejo.

— No me lo creo —dijo Elena—. Nadie escribe así.

     Me encogí de hombros.
"""


# ===========================================================================
# La plantilla de novela
# ===========================================================================
def test_la_novela_esta_en_el_catalogo():
    assert "novela" in plantillas.CATALOGO
    assert plantillas.ORDEN[0] == "novela"


def test_el_formato_de_dialogo_coincide_con_el_documento():
    f = plantillas.NOVELA.formato_diálogo
    assert f.sangria_izquierda_cm == 0.63
    assert f.sangria_francesa_cm == 0.63
    assert f.espacio_posterior_pt == 18
    assert f.interlineado == "minimo"


def test_el_formato_de_narrador_coincide_con_el_documento():
    f = plantillas.NOVELA.formato_parrafo
    assert f.primera_linea_cm == 1.25
    assert f.espacio_posterior_pt == 18
    assert f.alineacion == "justificada"


def test_todas_las_plantillas_usan_la_raya():
    """La raya de diálogo del castellano, U+2014. Ni guion corto ni nada
    que se le parezca."""
    for clave in plantillas.ORDEN:
        marca = plantillas.obtener(clave).marca_dialogo
        assert marca == "\u2014", f"{clave} usa {marca!r} en vez de la raya"


def test_ya_no_existe_el_campo_de_personaje():
    """Se eliminó del todo: era la vía por la que F3 metía el nombre."""
    assert not hasattr(plantillas.NOVELA, "pide_personaje")


def test_los_centimetros_se_traducen_a_espacios():
    """En el editor no hay centímetros: se aproximan con caracteres de una
    fuente monoespaciada, donde cada uno mide 0,254 cm."""
    assert plantillas.cm_a_caracteres(0.63) == 2
    assert plantillas.cm_a_caracteres(1.25) == 5
    assert plantillas.cm_a_caracteres(0) == 0


def test_la_sangria_del_editor_coincide_con_la_traduccion():
    assert plantillas.SANGRIA_NARRADOR == " " * 5
    assert plantillas.NOVELA.sangria_parrafo == plantillas.SANGRIA_NARRADOR


# ===========================================================================
# Detección de diálogos
# ===========================================================================
@pytest.mark.parametrize(
    "linea, es",
    [
        ("— No me lo creo.", True),
        ("– Con guion largo.", True),
        ("- Con guion corto.", True),
        ("     Aquella noche no dormí.", False),
        ("", False),
        ("   ", False),
    ],
)
def test_deteccion_de_dialogo(linea, es):
    """Se aceptan los tres guiones: al dictar o al pegar de otro sitio se
    cuelan el largo y el corto, y no se puede exigir la raya perfecta."""
    assert exportar_word.es_dialogo(linea) is es


# ===========================================================================
# Exportación a Word con las medidas exactas
# ===========================================================================
def _abrir(ruta):
    docx = pytest.importorskip("docx", reason="python-docx no instalado")
    return docx.Document(str(ruta))


def test_exportar_aplica_las_medidas_del_dialogo(tmp_path):
    pytest.importorskip("docx")
    ruta = tmp_path / "novela.docx"
    exportar_word.exportar(TEXTO, ruta, plantillas.NOVELA)

    documento = _abrir(ruta)
    dialogo = next(p for p in documento.paragraphs
                   if exportar_word.es_dialogo(p.text))
    f = dialogo.paragraph_format

    assert round(f.left_indent.cm, 2) == 0.63
    # La sangría francesa es una primera línea NEGATIVA: la raya sale al
    # margen y el resto del diálogo queda alineado debajo.
    assert round(f.first_line_indent.cm, 2) == -0.63
    assert f.space_after.pt == 18


def test_exportar_aplica_las_medidas_del_narrador(tmp_path):
    pytest.importorskip("docx")
    ruta = tmp_path / "novela.docx"
    exportar_word.exportar(TEXTO, ruta, plantillas.NOVELA)

    documento = _abrir(ruta)
    narrador = next(p for p in documento.paragraphs
                    if p.text.strip() and not exportar_word.es_dialogo(p.text))
    f = narrador.paragraph_format

    assert round(f.first_line_indent.cm, 2) == 1.25
    assert f.space_after.pt == 18
    assert f.left_indent is None or f.left_indent.cm == 0


def test_el_narrador_va_justificado(tmp_path):
    pytest.importorskip("docx")
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ruta = tmp_path / "novela.docx"
    exportar_word.exportar(TEXTO, ruta, plantillas.NOVELA)

    narrador = next(p for p in _abrir(ruta).paragraphs
                    if p.text.strip() and not exportar_word.es_dialogo(p.text))
    assert narrador.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY


def test_el_interlineado_minimo_es_at_least(tmp_path):
    """'Mínimo' en Word es AT_LEAST, no EXACTLY. Con 'exacto' se recortarían
    las tildes y las mayúsculas acentuadas."""
    pytest.importorskip("docx")
    from docx.enum.text import WD_LINE_SPACING

    ruta = tmp_path / "novela.docx"
    exportar_word.exportar(TEXTO, ruta, plantillas.NOVELA)

    parrafo = next(p for p in _abrir(ruta).paragraphs if p.text.strip())
    assert parrafo.paragraph_format.line_spacing_rule == WD_LINE_SPACING.AT_LEAST


def test_las_lineas_en_blanco_no_se_convierten_en_parrafos(tmp_path):
    """En Word la separación la da el espacio posterior de 18 pt. Si además
    se dejaran las líneas vacías, la separación sería el doble."""
    pytest.importorskip("docx")
    ruta = tmp_path / "novela.docx"
    exportar_word.exportar(TEXTO, ruta, plantillas.NOVELA)

    parrafos = [p.text for p in _abrir(ruta).paragraphs]
    assert all(p.strip() for p in parrafos)
    assert len(parrafos) == 3


def test_exportar_con_titulo(tmp_path):
    pytest.importorskip("docx")
    ruta = tmp_path / "novela.docx"
    exportar_word.exportar(TEXTO, ruta, plantillas.NOVELA, titulo="Capítulo 1")
    assert _abrir(ruta).paragraphs[0].text == "Capítulo 1"


def test_exportar_crea_la_carpeta(tmp_path):
    pytest.importorskip("docx")
    ruta = tmp_path / "nueva" / "carpeta" / "novela.docx"
    exportar_word.exportar(TEXTO, ruta, plantillas.NOVELA)
    assert ruta.exists()


def test_describir_formato_menciona_las_medidas():
    frase = exportar_word.describir_formato(plantillas.NOVELA)
    assert "0.63" in frase
    assert "1.25" in frase
    assert "18" in frase


# ===========================================================================
# El perfil de Julián
# ===========================================================================
def test_el_perfil_trae_los_ajustes_de_julian():
    p = perfiles.perfil_julian()
    assert p["tema"] == "alto_contraste"
    assert p["plantilla"] == "novela"
    assert p["tamano_fuente"] == 20
    assert p["velocidad"] == 2
    assert p["cursor_parpadea"] is False


def test_el_perfil_hereda_lo_que_no_toca():
    """Solo lleva lo que se aparta de los valores por defecto: cuando se
    añada una opción nueva, el perfil la recibe sin tocarlo."""
    p = perfiles.perfil_julian()
    assert len(p["atajos"]) == 39
    assert "importar" in p["atajos"]
    assert p["max_caracteres"] == 200000


def test_los_metadatos_no_son_ajustes():
    p = perfiles.perfil_julian()
    assert "_nombre" not in p
    assert "_descripcion" not in p


def test_describir_el_perfil():
    frase = perfiles.describir(perfiles.PERFIL_JULIAN)
    assert "Julián" in frase
    assert "novela" in frase


# ===========================================================================
# Exportar e importar configuración
# ===========================================================================
def test_ida_y_vuelta_de_un_perfil(tmp_path):
    from vozclip import config

    ajustes = perfiles.perfil_julian()
    ajustes["velocidad"] = 5
    ajustes["tema"] = "claro"

    ruta = perfiles.exportar(ajustes, tmp_path / "mio.json", nombre="El mío")
    recuperado = perfiles.importar(ruta, base=config.DEFAULTS)

    assert recuperado["velocidad"] == 5
    assert recuperado["tema"] == "claro"
    assert recuperado["plantilla"] == "novela"


def test_exportar_sin_ruta_inventa_el_nombre(tmp_path, monkeypatch):
    """Quien no ve la pantalla no debería pelearse con un diálogo de guardar
    para algo tan sencillo."""
    monkeypatch.setattr(perfiles, "carpeta_perfiles", lambda: tmp_path)
    ruta = perfiles.exportar(perfiles.perfil_julian())
    assert ruta.exists()
    assert ruta.suffix == ".json"


def test_importar_un_perfil_inexistente(tmp_path):
    with pytest.raises(ErrorFuente, match="No encuentro"):
        perfiles.importar(tmp_path / "fantasma.json")


def test_importar_un_json_roto(tmp_path):
    ruta = tmp_path / "roto.json"
    ruta.write_text("{ esto no es json", encoding="utf-8")
    with pytest.raises(ErrorFuente, match="mal escrito"):
        perfiles.importar(ruta)


def test_importar_algo_que_no_es_un_perfil(tmp_path):
    ruta = tmp_path / "lista.json"
    ruta.write_text("[1, 2, 3]", encoding="utf-8")
    with pytest.raises(ErrorFuente, match="no contiene un perfil"):
        perfiles.importar(ruta)


def test_un_perfil_con_valores_imposibles_se_rechaza(tmp_path):
    """Aplicarlo dejaría el programa en un estado del que no se puede salir
    sin ver la pantalla."""
    import json

    ruta = tmp_path / "malo.json"
    ruta.write_text(json.dumps({"tema": "inventado"}), encoding="utf-8")
    with pytest.raises(ErrorFuente, match="no existe"):
        perfiles.importar(ruta)


# ===========================================================================
# Validación
# ===========================================================================
def test_un_perfil_vacio_es_valido():
    """Un perfil parcial es lo normal: solo lleva lo que cambia."""
    assert perfiles.validar({}) == []


def test_se_detecta_un_tema_inexistente():
    assert "tema" in perfiles.validar({"tema": "arcoiris"})[0]


def test_se_detecta_una_plantilla_inexistente():
    assert "plantilla" in perfiles.validar({"plantilla": "haiku"})[0]


def test_se_detecta_un_tamano_fuera_de_rango():
    assert perfiles.validar({"tamano_fuente": 200})
    assert perfiles.validar({"tamano_fuente": 1})
    assert perfiles.validar({"tamano_fuente": "grande"})


def test_se_detecta_una_velocidad_fuera_de_rango():
    assert perfiles.validar({"velocidad": 99})
    assert perfiles.validar({"velocidad": -99})


def test_se_detecta_un_modo_inexistente():
    assert perfiles.validar({"modo": "telepatico"})


def test_los_valores_correctos_pasan():
    assert perfiles.validar({
        "tema": "alto_contraste",
        "plantilla": "novela",
        "tamano_fuente": 20,
        "velocidad": 2,
        "volumen": 100,
        "modo": "editor",
    }) == []
