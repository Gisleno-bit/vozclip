"""Tests de importación y exportación de archivos.

La regla que más importa aquí: importar NO debe tocar las sangrías. Un guion
de cine importado con los márgenes destrozados es inservible para quien no
puede verlos para arreglarlos.
"""

from __future__ import annotations

import pytest

from vozclip import fuentes

GUION_CINE = """INT. TALLER - NOCHE

Una mesa y una lámpara.

                    ELENA
               (sin levantar la vista)
          No me lo creo.
"""


# ===========================================================================
# Texto plano
# ===========================================================================
def test_importar_conserva_las_sangrias(tmp_path):
    """Lo esencial de toda esta funcionalidad."""
    ruta = tmp_path / "guion.txt"
    ruta.write_text(GUION_CINE, encoding="utf-8")

    salida = fuentes.leer_para_editor(ruta)
    lineas = salida.split("\n")

    assert lineas[4].startswith(" " * 20), "El personaje perdió su margen"
    assert lineas[5].startswith(" " * 15), "La acotación perdió su margen"
    assert lineas[6].startswith(" " * 10), "El diálogo perdió su margen"


def test_importar_conserva_las_lineas_en_blanco(tmp_path):
    ruta = tmp_path / "guion.txt"
    ruta.write_text(GUION_CINE, encoding="utf-8")
    assert "\n\n" in fuentes.leer_para_editor(ruta)


def test_el_camino_de_la_voz_si_normaliza(tmp_path):
    """Contraste deliberado entre los dos caminos.

    Importar conserva las sangrías intactas. El camino de la voz pasa por
    `texto.limpiar`, que junta líneas y quita márgenes para que se lea de
    corrido. Los dos son correctos; lo que sería un error es usar el
    segundo para importar.
    """
    from vozclip import texto as textoutil

    ruta = tmp_path / "guion.txt"
    ruta.write_text(GUION_CINE, encoding="utf-8")

    para_editar = fuentes.leer_para_editor(ruta)
    para_hablar = textoutil.limpiar(para_editar)

    assert "                    ELENA" in para_editar
    assert "                    ELENA" not in para_hablar


def test_importar_normaliza_los_saltos_de_windows(tmp_path):
    ruta = tmp_path / "windows.txt"
    ruta.write_bytes(b"una\r\ndos\r\ntres")
    salida = fuentes.leer_para_editor(ruta)
    assert "\r" not in salida
    assert salida.count("\n") == 2


def test_importar_texto_antiguo_en_windows_1252(tmp_path):
    ruta = tmp_path / "viejo.txt"
    ruta.write_bytes("Año de gestión".encode("cp1252"))
    assert "Año" in fuentes.leer_para_editor(ruta)


def test_importar_markdown(tmp_path):
    ruta = tmp_path / "notas.md"
    ruta.write_text("# Título\n\n  - Punto sangrado", encoding="utf-8")
    salida = fuentes.leer_para_editor(ruta)
    assert "# Título" in salida
    assert "  - Punto sangrado" in salida   # la viñeta NO se limpia al editar


# ===========================================================================
# Errores
# ===========================================================================
def test_archivo_inexistente(tmp_path):
    with pytest.raises(fuentes.ErrorFuente, match="No encuentro"):
        fuentes.leer_para_editor(tmp_path / "fantasma.txt")


def test_archivo_vacio(tmp_path):
    ruta = tmp_path / "vacio.txt"
    ruta.write_text("   \n  ", encoding="utf-8")
    with pytest.raises(fuentes.ErrorFuente, match="vacío"):
        fuentes.leer_para_editor(ruta)


def test_un_pdf_no_se_puede_editar(tmp_path):
    """Se explica por qué, en vez de dar un error genérico."""
    ruta = tmp_path / "documento.pdf"
    ruta.write_bytes(b"%PDF-1.4 falso")
    with pytest.raises(fuentes.ErrorFuente, match="no se pueden editar"):
        fuentes.leer_para_editor(ruta)


def test_archivo_demasiado_grande(tmp_path):
    ruta = tmp_path / "enorme.txt"
    ruta.write_text("a" * 5000, encoding="utf-8")
    with pytest.raises(fuentes.ErrorFuente, match="demasiado grande"):
        fuentes.leer_para_editor(ruta, max_caracteres=1000)


def test_una_carpeta_no_es_un_archivo(tmp_path):
    with pytest.raises(fuentes.ErrorFuente, match="no es un archivo"):
        fuentes.leer_para_editor(tmp_path)


# ===========================================================================
# RTF, con lector propio para no añadir dependencias
# ===========================================================================
def _escribir_rtf(ruta, cuerpo: str) -> None:
    ruta.write_text(
        r"{\rtf1\ansi\deff0{\fonttbl{\f0\fnil Courier;}}" + cuerpo + "}",
        encoding="latin-1",
    )


def test_rtf_basico(tmp_path):
    ruta = tmp_path / "texto.rtf"
    _escribir_rtf(ruta, r"\pard Primera linea\par Segunda linea\par")
    salida = fuentes.leer_para_editor(ruta)
    assert "Primera linea" in salida
    assert "Segunda linea" in salida


def test_rtf_con_acentos_escapados(tmp_path):
    ruta = tmp_path / "acentos.rtf"
    _escribir_rtf(ruta, r"\pard Ma\'f1ana ser\'e1 otro d\'eda\par")
    salida = fuentes.leer_para_editor(ruta)
    assert "Mañana" in salida
    assert "será" in salida


def test_rtf_convierte_tabuladores_en_sangria(tmp_path):
    ruta = tmp_path / "sangrado.rtf"
    _escribir_rtf(ruta, r"\pard\tab\tab ELENA\par")
    salida = fuentes.leer_para_editor(ruta)
    assert "        ELENA" in salida   # dos tabuladores = ocho espacios


def test_rtf_descarta_las_tablas_de_fuentes(tmp_path):
    """La tabla de fuentes no es texto del documento."""
    ruta = tmp_path / "fuentes.rtf"
    _escribir_rtf(ruta, r"\pard Texto de verdad\par")
    salida = fuentes.leer_para_editor(ruta)
    assert "Courier" not in salida
    assert "fnil" not in salida


def test_rtf_con_llaves_escapadas(tmp_path):
    ruta = tmp_path / "llaves.rtf"
    _escribir_rtf(ruta, r"\pard Una llave \{ y otra \}\par")
    salida = fuentes.leer_para_editor(ruta)
    assert "{" in salida and "}" in salida


def test_rtf_unicode(tmp_path):
    ruta = tmp_path / "unicode.rtf"
    _escribir_rtf(ruta, r"\pard Raya \u8212 ? de dialogo\par")
    assert "—" in fuentes.leer_para_editor(ruta)


# ===========================================================================
# Word
# ===========================================================================
def test_importar_docx(tmp_path):
    docx = pytest.importorskip("docx", reason="python-docx no instalado")
    ruta = tmp_path / "guion.docx"
    documento = docx.Document()
    documento.add_paragraph("INT. CASA - DÍA")
    documento.add_paragraph("Entra Elena.")
    documento.save(str(ruta))

    salida = fuentes.leer_para_editor(ruta)
    assert "INT. CASA - DÍA" in salida
    assert "Entra Elena." in salida


def test_docx_traduce_la_sangria_de_parrafo(tmp_path):
    """En Word la sangría está en el estilo, no en el texto. Si no se
    traduce a espacios, el guion importado pierde sus márgenes."""
    docx = pytest.importorskip("docx", reason="python-docx no instalado")
    from docx.shared import Inches

    ruta = tmp_path / "sangrado.docx"
    documento = docx.Document()
    parrafo = documento.add_paragraph("ELENA")
    parrafo.paragraph_format.left_indent = Inches(2.0)
    documento.save(str(ruta))

    salida = fuentes.leer_para_editor(ruta)
    assert salida.startswith(" " * 15), f"Sangría perdida: {salida!r}"


# ===========================================================================
# Guardado
# ===========================================================================
def test_guardar_texto(tmp_path):
    ruta = tmp_path / "salida.txt"
    fuentes.guardar_texto(ruta, GUION_CINE)
    assert ruta.read_text(encoding="utf-8") == GUION_CINE


def test_guardar_crea_la_carpeta(tmp_path):
    ruta = tmp_path / "nueva" / "carpeta" / "guion.txt"
    fuentes.guardar_texto(ruta, "contenido")
    assert ruta.exists()


def test_guardar_no_deja_temporales(tmp_path):
    ruta = tmp_path / "guion.txt"
    fuentes.guardar_texto(ruta, "contenido")
    assert list(tmp_path.glob("*.tmp")) == []


def test_ida_y_vuelta_conserva_el_formato(tmp_path):
    """Guardar y volver a importar debe devolver exactamente lo mismo."""
    ruta = tmp_path / "guion.txt"
    fuentes.guardar_texto(ruta, GUION_CINE)
    assert fuentes.leer_para_editor(ruta) == GUION_CINE


def test_las_extensiones_importables_estan_declaradas():
    for extension in (".txt", ".md", ".rtf", ".docx"):
        assert extension in fuentes.EXTENSIONES_IMPORTABLES
