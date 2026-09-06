"""Exportación a Word con el formato exacto del documento de estilo.

=============================================================================
POR QUÉ ESTE MÓDULO EXISTE
=============================================================================
El documento de estilo de Julián pide cosas que el texto plano NO PUEDE
guardar:

    Diálogos:  sangría izquierda 0,63 cm con sangría francesa 0,63,
               espacio posterior 18 pt, interlineado mínimo.
    Narrador:  primera línea 1,25 cm, espacio posterior 18 pt, justificado.

Un .txt no tiene centímetros, ni puntos de espaciado, ni justificación. En
el editor de VozClip esas sangrías se aproximan con espacios, que sirve para
escribir y para orientarse de oído, pero no es el formato de verdad.

Aquí sí lo es. Se construye un .docx de Word aplicando las medidas exactas,
párrafo a párrafo, distinguiendo automáticamente los diálogos (empiezan por
raya) del texto de narrador.
=============================================================================
"""

from __future__ import annotations

from pathlib import Path

from .fuentes import ErrorFuente
from .plantillas import FormatoWord, Plantilla

# Caracteres con los que puede empezar una línea de diálogo. La raya (—) es
# la correcta en castellano, pero al dictar o al pegar desde otro sitio se
# cuelan el guion largo y el corto, así que se aceptan los tres.
MARCAS_DIALOGO = ("—", "–", "-")


def es_dialogo(linea: str) -> bool:
    """¿Esta línea es una intervención hablada?"""
    limpia = linea.lstrip()
    return bool(limpia) and limpia[0] in MARCAS_DIALOGO


def exportar(
    texto: str,
    ruta: str | Path,
    plantilla: Plantilla,
    titulo: str | None = None,
) -> Path:
    """Escribe el documento en .docx con el formato de la plantilla.

    Cada párrafo se clasifica solo: si empieza por raya, se le aplica el
    formato de diálogo; si no, el de narrador.
    """
    try:
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.shared import Cm, Pt
    except ImportError as e:
        raise ErrorFuente(
            "Falta la librería python-docx para exportar a Word."
        ) from e

    ruta = Path(ruta).expanduser()
    ruta.parent.mkdir(parents=True, exist_ok=True)

    documento = docx.Document()

    if titulo:
        encabezado = documento.add_paragraph(titulo)
        encabezado.alignment = WD_ALIGN_PARAGRAPH.CENTER
        encabezado.paragraph_format.space_after = Pt(24)
        for trozo in encabezado.runs:
            trozo.bold = True

    formato_narrador = plantilla.formato_parrafo or FormatoWord()
    formato_dialogo = plantilla.formato_diálogo or formato_narrador

    for parrafo_texto in _partir_en_parrafos(texto):
        formato = formato_dialogo if es_dialogo(parrafo_texto) else formato_narrador
        parrafo = documento.add_paragraph(parrafo_texto.strip())
        _aplicar(parrafo.paragraph_format, formato,
                 Cm, Pt, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING)

    try:
        documento.save(str(ruta))
    except OSError as e:
        raise ErrorFuente(
            f"No he podido guardar {ruta.name}: {e.strerror or 'error de disco'}"
        ) from e

    return ruta


def _partir_en_parrafos(texto: str) -> list[str]:
    """Cada línea no vacía es un párrafo de Word.

    Las líneas en blanco del texto plano se descartan: en Word la separación
    entre párrafos la da el espacio posterior de 18 pt, no una línea vacía.
    Dejarlas produciría el doble de separación de la pedida.
    """
    return [linea for linea in texto.replace("\r\n", "\n").split("\n") if linea.strip()]


def _aplicar(destino, formato: FormatoWord, Cm, Pt, Alineacion, Interlineado) -> None:
    """Traduce un `FormatoWord` a las propiedades de python-docx."""
    if formato.sangria_izquierda_cm:
        destino.left_indent = Cm(formato.sangria_izquierda_cm)

    if formato.sangria_francesa_cm:
        # En Word, la sangría francesa es una primera línea NEGATIVA: el
        # texto va sangrado y la primera línea sale hacia el margen. Es lo
        # que hace que la raya quede fuera y el diálogo alineado debajo.
        destino.first_line_indent = Cm(-formato.sangria_francesa_cm)
    elif formato.primera_linea_cm:
        destino.first_line_indent = Cm(formato.primera_linea_cm)

    if formato.espacio_posterior_pt:
        destino.space_after = Pt(formato.espacio_posterior_pt)

    if formato.interlineado == "minimo":
        # "Mínimo" en Word = AT_LEAST: la línea nunca baja de ese alto, pero
        # crece si hace falta. No es lo mismo que "exacto", que recortaría
        # las tildes y los acentos altos.
        destino.line_spacing_rule = Interlineado.AT_LEAST

    if formato.alineacion == "justificada":
        destino.alignment = Alineacion.JUSTIFY


def describir_formato(plantilla: Plantilla) -> str:
    """Frase para decir en voz alta qué formato se va a aplicar."""
    narrador = plantilla.formato_parrafo
    dialogo = plantilla.formato_diálogo
    if narrador is None and dialogo is None:
        return "Sin formato especial."

    partes = []
    if dialogo and dialogo.sangria_francesa_cm:
        partes.append(f"diálogos con sangría francesa de {dialogo.sangria_francesa_cm}")
    if narrador and narrador.primera_linea_cm:
        partes.append(f"narrador con primera línea de {narrador.primera_linea_cm}")
    if narrador and narrador.espacio_posterior_pt:
        partes.append(f"{narrador.espacio_posterior_pt:.0f} puntos entre párrafos")
    return ", ".join(partes) + "." if partes else "Formato estándar."
