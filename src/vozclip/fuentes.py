"""De dónde sale el texto que hay que leer.

Tres fuentes:
  1. El portapapeles (lo que el usuario ha copiado con Ctrl+C).
  2. La selección actual (simulamos Ctrl+C y leemos el portapapeles,
     devolviéndolo después a su estado original).
  3. Un fichero del disco: .txt, .md, .docx o .pdf.
"""

from __future__ import annotations

import time
from pathlib import Path

# Extensiones que sabemos abrir
EXTENSIONES_TEXTO = {".txt", ".md", ".csv", ".log", ".json", ".xml", ".html", ".htm"}


class ErrorFuente(Exception):
    """Algo ha impedido obtener el texto. Siempre lleva un mensaje en
    castellano listo para leerse en voz alta."""


# ---------------------------------------------------------------------------
# Portapapeles
# ---------------------------------------------------------------------------
def leer_portapapeles() -> str:
    """Devuelve el contenido de texto del portapapeles."""
    import pyperclip  # import perezoso: los tests no lo necesitan

    try:
        return pyperclip.paste() or ""
    except Exception as e:  # pyperclip lanza excepciones muy variadas
        raise ErrorFuente("No he podido leer el portapapeles.") from e


def escribir_portapapeles(texto: str) -> None:
    import pyperclip

    try:
        pyperclip.copy(texto)
    except Exception:
        pass  # no es crítico: si falla, seguimos


def capturar_seleccion(espera: float = 0.25) -> str:
    """Copia lo que esté seleccionado en la ventana activa y lo devuelve.

    Truco clásico: guardamos el portapapeles, mandamos Ctrl+C, esperamos un
    momento a que la aplicación responda, leemos, y restauramos lo que había.
    Así no le destrozamos el portapapeles al usuario.
    """
    from pynput.keyboard import Controller, Key

    anterior = ""
    try:
        anterior = leer_portapapeles()
    except ErrorFuente:
        pass

    teclado = Controller()
    # Marcamos el portapapeles para detectar si la copia funcionó
    marca = "\x00vozclip\x00"
    escribir_portapapeles(marca)

    with teclado.pressed(Key.ctrl):
        teclado.press("c")
        teclado.release("c")

    time.sleep(espera)

    try:
        nuevo = leer_portapapeles()
    except ErrorFuente:
        nuevo = ""

    # Restauramos lo que el usuario tenía copiado
    escribir_portapapeles(anterior)

    if nuevo == marca or not nuevo.strip():
        raise ErrorFuente("No hay nada seleccionado, o la aplicación no permite copiar.")
    return nuevo


# ---------------------------------------------------------------------------
# Ficheros
# ---------------------------------------------------------------------------
def leer_fichero(ruta: str | Path, max_caracteres: int = 200000) -> str:
    """Extrae el texto de un fichero soportado."""
    ruta = Path(ruta).expanduser()

    if not ruta.exists():
        raise ErrorFuente(f"No encuentro el archivo {ruta.name}.")
    if not ruta.is_file():
        raise ErrorFuente(f"{ruta.name} no es un archivo.")

    sufijo = ruta.suffix.lower()

    if sufijo in EXTENSIONES_TEXTO:
        texto = _leer_texto_plano(ruta)
    elif sufijo == ".docx":
        texto = _leer_docx(ruta)
    elif sufijo == ".pdf":
        texto = _leer_pdf(ruta)
    else:
        raise ErrorFuente(
            f"No sé abrir archivos de tipo {sufijo or 'desconocido'}. "
            "Puedo con texto, docx y pdf."
        )

    if not texto.strip():
        raise ErrorFuente(f"El archivo {ruta.name} no contiene texto legible.")

    if len(texto) > max_caracteres:
        texto = texto[:max_caracteres] + " … Fin del fragmento. El archivo continúa."
    return texto


def _leer_texto_plano(ruta: Path) -> str:
    """Prueba varias codificaciones: los .txt españoles suelen venir en
    UTF-8, pero los antiguos vienen en Windows-1252."""
    for codificacion in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return ruta.read_text(encoding=codificacion)
        except (UnicodeDecodeError, LookupError):
            continue
    raise ErrorFuente(f"No he podido descifrar el texto de {ruta.name}.")


def _leer_docx(ruta: Path) -> str:
    try:
        import docx  # paquete python-docx
    except ImportError as e:
        raise ErrorFuente(
            "Falta la librería python-docx para leer documentos de Word."
        ) from e

    try:
        documento = docx.Document(str(ruta))
    except Exception as e:
        raise ErrorFuente(f"El archivo {ruta.name} no se ha podido abrir.") from e

    partes = [p.text for p in documento.paragraphs if p.text.strip()]

    # Las tablas también llevan texto útil (facturas, listados...)
    for tabla in documento.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells if c.text.strip()]
            if celdas:
                partes.append(", ".join(celdas))

    return "\n\n".join(partes)


def _leer_pdf(ruta: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ErrorFuente("Falta la librería pypdf para leer archivos PDF.") from e

    try:
        lector = PdfReader(str(ruta))
    except Exception as e:
        raise ErrorFuente(f"El archivo {ruta.name} no se ha podido abrir.") from e

    paginas = []
    for pagina in lector.pages:
        try:
            paginas.append(pagina.extract_text() or "")
        except Exception:
            continue

    texto = "\n\n".join(p for p in paginas if p.strip())
    if not texto.strip():
        raise ErrorFuente(
            f"{ruta.name} parece un PDF escaneado, sin texto. "
            "Haría falta reconocimiento óptico de caracteres."
        )
    return texto


# ---------------------------------------------------------------------------
# Importación al editor propio
# ---------------------------------------------------------------------------
# Ojo a la diferencia con `leer_fichero`: aquel prepara texto para LEERLO en
# voz alta, y por eso lo normaliza. Este prepara texto para EDITARLO, así que
# conserva escrupulosamente los espacios del principio de cada línea. Si se
# normalizara, un guion de cine importado perdería sus márgenes y quedaría
# irreconocible para quien no puede verlos.
EXTENSIONES_IMPORTABLES = {".txt", ".md", ".rtf", ".docx", ".text", ".log", ".csv"}


def leer_para_editor(ruta: str | Path, max_caracteres: int = 2_000_000) -> str:
    """Carga un archivo conservando sangrías y saltos de línea tal cual."""
    ruta = Path(ruta).expanduser()

    if not ruta.exists():
        raise ErrorFuente(f"No encuentro el archivo {ruta.name}.")
    if not ruta.is_file():
        raise ErrorFuente(f"{ruta.name} no es un archivo.")

    sufijo = ruta.suffix.lower()

    if sufijo == ".rtf":
        texto = _leer_rtf(ruta)
    elif sufijo == ".docx":
        texto = _leer_docx_para_editor(ruta)
    elif sufijo == ".pdf":
        raise ErrorFuente(
            "Los PDF no se pueden editar. Puedo leerlo en voz alta, pero no "
            "abrirlo para escribir."
        )
    else:
        texto = _leer_texto_plano(ruta)

    texto = texto.replace("\r\n", "\n").replace("\r", "\n")

    if not texto.strip():
        raise ErrorFuente(f"El archivo {ruta.name} está vacío.")
    if len(texto) > max_caracteres:
        raise ErrorFuente(
            f"{ruta.name} es demasiado grande para editarlo aquí."
        )
    return texto


def _leer_docx_para_editor(ruta: Path) -> str:
    """Como `_leer_docx`, pero conservando el formato para editar.

    La sangría de un .docx no está en el texto sino en el estilo del
    párrafo, así que se traduce a espacios: es la única forma de que un
    guion importado desde Word conserve sus márgenes en un editor de texto
    plano.
    """
    try:
        import docx
    except ImportError as e:
        raise ErrorFuente(
            "Falta la librería python-docx para abrir documentos de Word."
        ) from e

    try:
        documento = docx.Document(str(ruta))
    except Exception as e:
        raise ErrorFuente(f"El archivo {ruta.name} no se ha podido abrir.") from e

    lineas = []
    for parrafo in documento.paragraphs:
        sangria = ""
        try:
            izquierda = parrafo.paragraph_format.left_indent
            if izquierda is not None:
                # Un carácter monoespaciado ocupa aproximadamente 1/10 de
                # pulgada a 12 puntos. 914400 EMU = 1 pulgada.
                espacios = int(izquierda.emu / 914400 * 10)
                sangria = " " * max(0, min(60, espacios))
        except Exception:
            sangria = ""
        lineas.append(sangria + parrafo.text)

    return "\n".join(lineas)


def _leer_rtf(ruta: Path) -> str:
    """Extrae el texto de un RTF sin instalar nada.

    Un RTF es texto plano con órdenes que empiezan por barra invertida. Este
    lector cubre lo que hace falta para un guion: párrafos, tabulaciones,
    acentos escapados y grupos que hay que descartar (fuentes, colores).
    No pretende ser un intérprete completo de RTF; para eso haría falta una
    librería, y no compensa cargar el ejecutable por un formato de paso.
    """
    crudo = ruta.read_text(encoding="latin-1", errors="replace")

    # Grupos de control que no contienen texto visible
    for bloque in ("fonttbl", "colortbl", "stylesheet", "info", "pict"):
        crudo = _quitar_grupo(crudo, bloque)

    salida: list[str] = []
    i = 0
    largo = len(crudo)

    while i < largo:
        caracter = crudo[i]

        if caracter == "\\":
            i += 1
            if i >= largo:
                break
            siguiente = crudo[i]

            # Carácter escapado: \{ \} \\
            if siguiente in "{}\\":
                salida.append(siguiente)
                i += 1
                continue

            # Byte en hexadecimal: \'e1 -> á
            if siguiente == "'" and i + 2 < largo:
                try:
                    salida.append(bytes([int(crudo[i + 1 : i + 3], 16)]).decode("latin-1"))
                except ValueError:
                    pass
                i += 3
                continue

            # Palabra de control: \par, \tab, \u241...
            palabra = ""
            while i < largo and (crudo[i].isalpha()):
                palabra += crudo[i]
                i += 1
            parametro = ""
            if i < largo and (crudo[i] == "-" or crudo[i].isdigit()):
                while i < largo and (crudo[i] == "-" or crudo[i].isdigit()):
                    parametro += crudo[i]
                    i += 1
            if i < largo and crudo[i] == " ":
                i += 1   # el espacio tras una orden es un separador, no texto

            if palabra in ("par", "line", "pard"):
                if palabra != "pard":
                    salida.append("\n")
            elif palabra == "tab":
                salida.append("    ")
            elif palabra == "u" and parametro:
                try:
                    salida.append(chr(int(parametro) % 65536))
                except ValueError:
                    pass
            continue

        if caracter in "{}":
            i += 1
            continue

        if caracter in "\r\n":
            i += 1
            continue

        salida.append(caracter)
        i += 1

    return "".join(salida)


def _quitar_grupo(texto: str, nombre: str) -> str:
    """Elimina un grupo RTF completo, contando llaves para no pasarse."""
    marca = "{\\" + nombre
    while True:
        inicio = texto.find(marca)
        if inicio == -1:
            return texto
        profundidad = 0
        i = inicio
        while i < len(texto):
            if texto[i] == "{":
                profundidad += 1
            elif texto[i] == "}":
                profundidad -= 1
                if profundidad == 0:
                    break
            i += 1
        texto = texto[:inicio] + texto[i + 1 :]


def guardar_texto(ruta: str | Path, contenido: str) -> Path:
    """Escribe el documento en disco de forma atómica.

    Primero a un temporal y luego se renombra: si se corta la luz a mitad,
    el archivo anterior sigue intacto en vez de quedar a medias.
    """
    ruta = Path(ruta).expanduser()
    ruta.parent.mkdir(parents=True, exist_ok=True)

    temporal = ruta.with_suffix(ruta.suffix + ".tmp")
    try:
        temporal.write_text(contenido, encoding="utf-8")
        temporal.replace(ruta)
    except OSError as e:
        raise ErrorFuente(
            f"No he podido guardar {ruta.name}: {e.strerror or 'error de disco'}"
        ) from e
    return ruta
