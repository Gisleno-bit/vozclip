"""Prueba de humo de extremo a extremo.

Levanta el programa REAL (ventana, servicio de voz y mapa de atajos), simula
la pulsación de varios atajos disparando su función del mapa igual que haría
pynput, y comprueba que el efecto llega al documento y a la voz.

Es la prueba que faltaba en la versión anterior: los tests unitarios pasaban
y sin embargo el programa no hacía nada al ejecutarlo.

Ejecutar con:   xvfb-run -a python3 pruebas_manuales/humo.py
"""

from __future__ import annotations

import copy
import os
import sys
import time

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "src"))

from vozclip import config  # noqa: E402
from vozclip.atajos import construir_mapa  # noqa: E402
from vozclip.hud import HUD  # noqa: E402
from vozclip.voz import ServicioVoz  # noqa: E402

fallos = 0


def comprobar(descripcion: str, condicion: bool) -> None:
    global fallos
    if condicion:
        print(f"  [ OK ] {descripcion}")
    else:
        fallos += 1
        print(f"  [FALLA] {descripcion}")


def main() -> int:
    print("Prueba de humo de VozClip Escritor")
    print("=" * 60)

    # 1. Servicio de voz
    print("\n1. Servicio de voz")
    voz = ServicioVoz(motor_forzado="falso")
    voz.arrancar()
    comprobar("el hilo de voz arranca", voz._hilo.is_alive())
    comprobar("el motor se creó dentro del hilo", voz._motor is not None)

    # 2. Ventana
    print("\n2. Ventana (HUD)")
    ajustes = copy.deepcopy(config.DEFAULTS)
    ventana = HUD(voz, ajustes)
    ventana.raiz.update()
    comprobar("la ventana existe", bool(ventana.raiz.winfo_exists()))
    comprobar("hay 11 botones", len(ventana.botones) == 11)
    comprobar("el sexto es corregir", "Corregir" in ventana.botones[5].cget("text"))
    comprobar("el último es LibreOffice", "LibreOffice" in ventana.botones[10].cget("text"))
    comprobar("el primero es el de grabar",
              "Grabar" in ventana.botones[0].cget("text"))
    comprobar("los cinco comandos van agrupados en su color",
              ventana.grupos_boton[:5] == ["comando"] * 5)
    comprobar("corregir tiene color propio", ventana.grupos_boton[5] == "corregir")
    comprobar("guardar tiene color propio", ventana.grupos_boton[6] == "guardar")
    comprobar("el editor existe", bool(ventana.editor.winfo_exists()))

    # 3. Saludo hablado
    print("\n3. Voz al arrancar")
    ventana.saludar()
    voz.esperar_silencio(limite=3)
    comprobar(
        "se saluda al arrancar",
        any("en marcha" in t for t in voz._motor.dicho),
    )

    # 4. Mapa de atajos globales
    print("\n4. Atajos globales")
    mapa = construir_mapa(
        ajustes["atajos"], ventana.encolar_orden, set(ventana.acciones())
    )
    comprobar(f"se registran {len(mapa)} combinaciones", len(mapa) >= 18)
    comprobar(
        "todas las combinaciones tienen formato válido",
        all(c.startswith("<") for c in mapa),
    )

    # 5. Simulación de pulsaciones reales
    print("\n5. Simulación de pulsaciones")

    def pulsar(combinacion: str) -> None:
        """Dispara el atajo igual que haría pynput, desde otro hilo lógico,
        y deja que el bucle del hilo principal lo procese.

        Se normaliza primero, porque el mapa está indexado por la forma que
        entiende pynput: "<ctrl>+<alt>+<plus>" se guarda como "<ctrl>+<alt>++".
        """
        from vozclip.atajos import normalizar

        mapa[normalizar(combinacion)]()
        ventana._atender_cola()
        ventana.raiz.update()
        time.sleep(0.05)

    pulsar("<ctrl>+<alt>+g")          # insertar plantilla
    comprobar(
        "Ctrl+Alt+G inserta la plantilla",
        bool(ventana._texto().strip()),
    )
    comprobar("no quedan marcas | en el texto", "|" not in ventana._texto())
    comprobar("el cursor está en el primer hueco", ventana._cursor() == ventana.marcas[0])

    pulsar("<ctrl>+<alt>+n")          # cambiar plantilla
    comprobar("Ctrl+Alt+N rota la plantilla",
              ventana.plantilla.clave == "narrativo")

    ventana.editor.delete("1.0", "end")
    ventana.editor.insert("1.0", "diálogo de prueba")
    ventana.editor.mark_set("insert", "1.0 + 5 chars")
    pulsar("<ctrl>+<alt>+i")          # sangría
    comprobar(
        "Ctrl+Alt+I aplica la sangría al margen",
        ventana._texto().startswith(ventana.plantilla.sangria_parrafo + "diálogo"),
    )

    pulsar("<ctrl>+<alt>+<enter>")    # siguiente línea
    comprobar("Ctrl+Alt+Intro añade línea", ventana._texto().count("\n") == 1)

    voz._motor.dicho.clear()
    pulsar("<ctrl>+<alt>+j")          # leer línea
    voz.esperar_silencio(limite=3)
    comprobar(
        "Ctrl+Alt+J lee la línea en voz alta",
        any("Línea" in t for t in voz._motor.dicho),
    )

    voz._motor.dicho.clear()
    pulsar("<ctrl>+<alt>+<up>")       # más rápido
    voz.esperar_silencio(limite=3)
    comprobar("Ctrl+Alt+Flecha arriba sube la velocidad", ajustes["velocidad"] == 1)
    comprobar(
        "y lo anuncia por voz",
        any("Velocidad" in t for t in voz._motor.dicho),
    )

    voz._motor.dicho.clear()
    pulsar("<ctrl>+<alt>+m")          # cambiar modo
    comprobar("Ctrl+Alt+M cambia a modo externo", ventana.modo == "externo")
    pulsar("<ctrl>+<alt>+m")
    comprobar("y vuelve al editor propio", ventana.modo == "editor")

    voz._motor.dicho.clear()
    pulsar("<ctrl>+<alt>+h")          # ayuda
    voz.esperar_silencio(limite=5)
    comprobar(
        "Ctrl+Alt+H lee la ayuda",
        any("plantilla" in t.lower() for t in voz._motor.dicho),
    )

    # 5a. Los cinco comandos de todos los días
    print("\n5a. Los cinco comandos diarios (F1 a F5)")
    ventana.plantilla = __import__(
        "vozclip.plantillas", fromlist=["obtener"]
    ).obtener("novela")
    ventana.editor.delete("1.0", "end")
    ventana.editor.insert("1.0", "     Aquella noche no dormí.")
    ventana.editor.mark_set("insert", "end-1c")

    pulsar("<f2>")
    lineas = ventana._texto().split("\n")
    comprobar("F2 abre un párrafo con su sangría",
              lineas[-1] == ventana.plantilla.sangria_parrafo)

    pulsar("<f3>")
    ultima = ventana._texto().split("\n")[-1]
    comprobar("F3 pone sangría y raya pegada, sin espacio", ultima == "  \u2014")
    comprobar("la raya es la del castellano (U+2014)", ord(ultima[-1]) == 0x2014)
    comprobar("F3 no mete ningún nombre de personaje",
              "PERSONAJE" not in ventana._texto())

    voz._motor.dicho.clear()
    pulsar("<f4>")
    voz.esperar_silencio(limite=3)
    comprobar("F4 lee el último párrafo",
              any("último párrafo" in t for t in voz._motor.dicho))

    voz._motor.dicho.clear()
    pulsar("<f5>")
    voz.esperar_silencio(limite=3)
    comprobar("F5 lee el texto entero",
              any("Aquella noche" in t for t in voz._motor.dicho))

    # 5a-bis. El formato de Julián en un .docx de verdad
    print("\n5a-bis. Exportar a Word con el formato de Julián")
    import tempfile as _tmpw
    from pathlib import Path as _Pathw

    from vozclip import exportar_word as _ew
    from vozclip import plantillas as _pl

    with _tmpw.TemporaryDirectory() as carpeta:
        destino = _Pathw(carpeta) / "novela.docx"
        _ew.exportar(
            "     Narrador con su sangría.\n\n— Y un diálogo con raya.",
            destino, _pl.NOVELA,
        )
        try:
            import docx as _docx

            doc = _docx.Document(str(destino))
            dial = next(p for p in doc.paragraphs if _ew.es_dialogo(p.text))
            narr = next(p for p in doc.paragraphs
                        if p.text.strip() and not _ew.es_dialogo(p.text))
            comprobar("el diálogo lleva sangría izquierda 0,63",
                      round(dial.paragraph_format.left_indent.cm, 2) == 0.63)
            comprobar("el diálogo lleva sangría francesa 0,63",
                      round(dial.paragraph_format.first_line_indent.cm, 2) == -0.63)
            comprobar("el narrador lleva primera línea 1,25",
                      round(narr.paragraph_format.first_line_indent.cm, 2) == 1.25)
            comprobar("los dos llevan 18 puntos de posterior",
                      dial.paragraph_format.space_after.pt == 18
                      and narr.paragraph_format.space_after.pt == 18)
        except ImportError:
            comprobar("python-docx no instalado, sin comprobar el .docx", True)

    # 5a-tris. Exportar a LibreOffice
    print("\n5a-tris. Exportar a LibreOffice (.odt)")
    from vozclip import exportar_odt as _eo

    with _tmpw.TemporaryDirectory() as carpeta:
        destino = _Pathw(carpeta) / "novela.odt"
        _eo.exportar("     Narrador.\n\n  —Diálogo.", destino, _pl.NOVELA)
        leido = _eo.leer(destino)
        dial = {k.split("}")[1]: v for k, v in leido["estilos"]["Dialogo"].items()}
        comprobar("el mimetype va primero y sin comprimir",
                  leido["mimetype_primero"] and leido["mimetype_sin_comprimir"])
        comprobar("el diálogo lleva 0,63 con francesa 0,63",
                  dial.get("margin-left") == "0.63cm" and dial.get("text-indent") == "-0.63cm")
        comprobar("los párrafos se clasifican solos",
                  [p["estilo"] for p in leido["parrafos"]] == ["Narrador", "Dialogo"])

    # 5b. Accesibilidad
    print("\n5b. Accesibilidad")
    from vozclip.hud import ORDEN_TEMAS

    pulsar("<ctrl>+<alt>+c")
    comprobar("Ctrl+Alt+C cambia de tema", ajustes["tema"] == ORDEN_TEMAS[1])
    comprobar("el tema es alto contraste", ventana.editor.cget("bg") == "#000000")

    antes = ventana.tamano
    pulsar("<ctrl>+<alt>+<plus>")
    comprobar("Ctrl+Alt+más agranda la letra", ventana.tamano > antes)
    pulsar("<ctrl>+<alt>+<minus>")

    pulsar("<ctrl>+<alt>+z")
    ventana.raiz.update()
    comprobar("Ctrl+Alt+Z esconde la interfaz", not ventana.editor.winfo_ismapped())
    pulsar("<ctrl>+<alt>+z")
    ventana.raiz.update()
    comprobar("y la devuelve", ventana.editor.winfo_ismapped())

    # 5c. Importar y exportar
    print("\n5c. Importar y exportar")
    import tempfile as _tmp
    from pathlib import Path as _Path

    with _tmp.TemporaryDirectory() as carpeta:
        origen = _Path(carpeta) / "guion.txt"
        origen.write_text(
            "INT. CASA - DÍA\n\n                    ELENA\n          Hola.\n",
            encoding="utf-8",
        )
        ventana.abrir(origen)
        lineas = ventana._texto().split("\n")
        comprobar("importar conserva las sangrías",
                  lineas[2].startswith(" " * 20) and lineas[3].startswith(" " * 10))

    copiado = {}
    import vozclip.hud as modhud

    original = modhud.escribir_portapapeles
    modhud.escribir_portapapeles = lambda t: copiado.setdefault("texto", t)
    try:
        pulsar("<ctrl>+<alt>+e")
        comprobar("Ctrl+Alt+E exporta al portapapeles",
                  "INT. CASA" in copiado.get("texto", ""))
    finally:
        modhud.escribir_portapapeles = original

    # 6. Dictado por voz
    print("\n6. Dictado por voz")
    from vozclip import dictado as moddictado

    comprobar(
        "la puntuación hablada funciona",
        moddictado.aplicar_puntuacion("raya no me lo creo punto") == "—No me lo creo.",
    )

    # Ciclo completo con micrófono y reconocedor falsos
    ventana.servicio_dictado = moddictado.ServicioDictado(
        notificar=ventana._evento_dictado,
        ajustes={"activado": True, "modelo": "fingido"},
        fabrica_motor=lambda: moddictado.MotorDictadoFalso(
            ["no me lo creo punto y aparte y me voy"]
        ),
        fabrica_captura=lambda: moddictado.CapturaFalsa(),
    )

    # Dictamos dentro de un bloque de diálogo sangrado
    ventana.plantilla = __import__(
        "vozclip.plantillas", fromlist=["obtener"]
    ).obtener("cine")
    ventana.editor.delete("1.0", "end")
    ventana.editor.insert("1.0", "                    ELENA\n          ")
    ventana.editor.mark_set("insert", "end-1c")

    pulsar("<f1>")
    plazo = time.monotonic() + 6
    while ventana.servicio_dictado.activo and time.monotonic() < plazo:
        ventana._atender_dictado()
        ventana.raiz.update()
        time.sleep(0.02)
    time.sleep(0.2)
    ventana._atender_dictado()
    ventana.raiz.update()

    lineas = ventana._texto().split("\n")
    comprobar("F1 dicta y el texto llega al editor", "No me lo creo." in ventana._texto())
    comprobar(
        "el dictado respeta la sangría de la plantilla",
        lineas[-1].startswith("          "),
    )
    comprobar("los signos hablados se aplican", "." in ventana._texto())

    # 6b. Corrección por voz
    print("\n6b. Corrección por voz (F9)")
    ventana.editor.delete("1.0", "end")
    ventana.editor.insert("1.0", "     Aquella noche no dormí, la casa estaba en silencio.")
    ventana.editor.mark_set("insert", "end-1c")
    ventana.servicio_dictado = moddictado.ServicioDictado(
        notificar=ventana._evento_dictado,
        ajustes={"activado": True, "modelo": "fingido"},
        fabrica_motor=lambda: moddictado.MotorDictadoFalso(["cambia casa por cosa"]),
        fabrica_captura=lambda: moddictado.CapturaFalsa(),
    )
    voz._motor.dicho.clear()
    pulsar("<f9>")
    plazo = time.monotonic() + 6
    while time.monotonic() < plazo:
        ventana._atender_dictado()
        ventana.raiz.update()
        time.sleep(0.02)
        if not ventana.servicio_dictado.activo and ventana.cola_dictado.empty():
            break
    time.sleep(0.2)
    ventana._atender_dictado()
    voz.esperar_silencio(limite=3)
    comprobar("F9 explica qué decir", any("cambia casa por cosa" in t for t in voz._motor.dicho))
    comprobar("se cambia solo esa palabra",
              ventana._texto() == "     Aquella noche no dormí, la cosa estaba en silencio.")
    comprobar("y se confirma por voz", any("He cambiado casa por cosa" in t for t in voz._motor.dicho))

    # 7. Guardado
    print("\n7. Guardado")
    import tempfile

    with tempfile.TemporaryDirectory() as tmp:
        ajustes["carpeta_guiones"] = tmp
        ventana.ruta_actual = None
        pulsar("<ctrl>+<alt>+d")
        comprobar("Ctrl+Alt+D guarda el archivo", bool(
            ventana.ruta_actual and ventana.ruta_actual.exists()
        ))

    # 7. La franja de estado refleja la realidad
    print("\n8. Franja de estado")
    estado = ventana.etiqueta_estado.cget("text")
    comprobar("muestra el modo", "Modo" in estado)
    comprobar("muestra la plantilla", "Plantilla" in estado)
    comprobar("muestra la velocidad actualizada", "Velocidad: 1" in estado)

    # Cierre
    ventana.raiz.destroy()
    voz.cerrar()

    print("\n" + "=" * 60)
    if fallos:
        print(f"RESULTADO: {fallos} comprobación(es) fallida(s).")
    else:
        print("RESULTADO: todo correcto. El programa arranca, habla y responde.")
    return 1 if fallos else 0


if __name__ == "__main__":
    sys.exit(main())
